import re
import unicodedata
from typing import Dict, List
from io import BytesIO
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.oxml.ns import qn


def normalizer(s: str) -> str:
    s = unicodedata.normalize("NFKD", s)
    s = "".join(c for c in s if not unicodedata.combining(c))
    s = s.lower()
    s = re.sub(r"[^a-z0-9]+", " ", s).strip()
    s = re.sub(r"\s+", " ", s)
    return s


def eliminar_linea(md: str) -> str:
    out = []
    for line in md.splitlines():
        if re.fullmatch(r"\s*-{3,}\s*", line):
            out.append("")
        else:
            out.append(line)
    return "\n".join(out)


def extraer_secciones_placeholders(archivo_md: str) -> Dict:
    with open(archivo_md, "r", encoding="utf-8") as f:
        texto = f.read()

    titulo_doc = None
    for ln in texto.splitlines():
        if ln.startswith("# "):
            titulo_doc = ln.replace("#", "").strip()
            break

    token_line = re.compile(r"(?m)^\s*\{\{\s*([^\}]+?)\s*\}\}\s*$")
    matches = list(token_line.finditer(texto))

    secciones: List[Dict] = []
    for i, m in enumerate(matches):
        key = m.group(1).strip()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(texto)
        contenido = texto[start:end].lstrip()
        contenido = eliminar_linea(contenido)
        secciones.append({"titulo": key, "contenido": contenido})

    return {"titulo_doc": titulo_doc, "secciones": secciones}


def _procesar_contenido_celda(contenido: str, tc_element):
    """
    Procesa contenido markdown en una celda de tabla, soportando negritas y listas.
    """
    contenido = contenido.strip()
    if not contenido:
        p = OxmlElement("w:p")
        tc_element.append(p)
        return

    # Dividir por líneas para manejar listas
    lineas = contenido.split('\n')
    
    for i, linea in enumerate(lineas):
        linea = linea.strip()
        if not linea:
            continue
            
        p = OxmlElement("w:p")
        
        # Detectar si es elemento de lista
        es_lista = False
        if linea.startswith('- '):
            linea = linea[2:].strip()
            es_lista = True
            # Agregar propiedades de lista al párrafo
            pPr = OxmlElement("w:pPr")
            numPr = OxmlElement("w:numPr")
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), "0")
            numId = OxmlElement("w:numId")
            numId.set(qn("w:val"), "1")
            numPr.append(ilvl)
            numPr.append(numId)
            pPr.append(numPr)
            p.append(pPr)
        
        # Procesar negritas en la línea
        parts = re.split(r"(\*\*[^*]+\*\*)", linea)
        
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                # Texto en negrita
                r = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rPr.append(b)
                r.append(rPr)
                t = OxmlElement("w:t")
                t.text = part[2:-2]
                r.append(t)
                p.append(r)
            elif part:
                # Texto normal
                r = OxmlElement("w:r")
                t = OxmlElement("w:t")
                t.text = part
                r.append(t)
                p.append(r)
        
        tc_element.append(p)


def render_tabla_md(md_lines: List[str], insert_idx: int, parent_element, doc) -> None:
    if len(md_lines) < 2:
        return
    header_raw = md_lines[0]
    separator_raw = md_lines[1]
    header = [h.strip() for h in header_raw.strip().strip("|").split("|")]
    separator_cells = [s.strip() for s in separator_raw.strip().strip("|").split("|")]
    if len(header) != len(separator_cells):
        return

    data_rows: List[List[str]] = []
    for line in md_lines[2:]:
        if not line.strip():
            continue
        if "|" not in line:
            break
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if len(cells) == len(header):
            data_rows.append(cells)

    def _get_first_section(container):
        # container puede ser Document o Body
        if hasattr(container, "sections"):
            try:
                return container.sections[0]
            except Exception:
                pass
        if hasattr(container, "part"):
            part = getattr(container, "part")
            doc_obj = getattr(part, "document", None)
            if doc_obj and hasattr(doc_obj, "sections"):
                try:
                    return doc_obj.sections[0]
                except Exception:
                    pass
        return None

    def _compute_col_widths(cols: List[str], rows: List[List[str]]) -> List[int]:
        max_len = [len(c) for c in cols]
        for r in rows:
            for i, c in enumerate(r):
                if len(c) > max_len[i]:
                    max_len[i] = len(c)
        total_chars = sum(max_len) or 1

        section = _get_first_section(doc)
        # Fallbacks si no hay sección disponible
        if section is not None:
            try:
                emu_to_twips = lambda emu: int(emu / 635)
                page_twips = emu_to_twips(section.page_width)
                left_twips = emu_to_twips(section.left_margin)
                right_twips = emu_to_twips(section.right_margin)
            except Exception:
                # Fallback si propiedades no accesibles
                page_twips = 12240
                left_twips = right_twips = 1440
        else:
            page_twips = 12240  # ~8.5in
            left_twips = right_twips = 1440  # ~1in cada margen

        avail = max(page_twips - left_twips - right_twips - 200, 3000)
        MIN_COL = 1500
        widths = [max(int(avail * l / total_chars), MIN_COL) for l in max_len]
        scale = avail / sum(widths)
        return [int(w * scale) for w in widths]

    col_widths = _compute_col_widths(header, data_rows)

    tbl = OxmlElement("w:tbl")
    parent_element.insert(insert_idx, tbl)

    # Agregar propiedades de tabla con bordes sólidos
    tblPr = OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    
    # Definir bordes sólidos para toda la tabla
    border_attrs = {
        qn("w:val"): "single",
        qn("w:sz"): "4",
        qn("w:space"): "0",
        qn("w:color"): "000000"
    }
    
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border_elem = OxmlElement(f"w:{border_name}")
        for attr, value in border_attrs.items():
            border_elem.set(attr, value)
        tblBorders.append(border_elem)
    
    tblPr.append(tblBorders)
    tbl.append(tblPr)

    grid = OxmlElement("w:tblGrid")
    for w in col_widths:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(w))
        grid.append(gridCol)
    tbl.append(grid)

    def add_row(values, bold=False):
        tr = OxmlElement("w:tr")
        for idx, val in enumerate(values):
            tc = OxmlElement("w:tc")
            tcPr = OxmlElement("w:tcPr")
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(col_widths[idx]))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
            
            # Agregar sombreado oscuro para headers
            if bold:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "404040")  # Gris oscuro
                tcPr.append(shd)
            
            tc.append(tcPr)
            
            if bold:
                # Para headers, texto blanco y negrita
                p = OxmlElement("w:p")
                r = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                color = OxmlElement("w:color")
                color.set(qn("w:val"), "FFFFFF")  # Texto blanco
                rPr.append(b)
                rPr.append(color)
                r.append(rPr)
                t = OxmlElement("w:t")
                t.text = val
                r.append(t)
                p.append(r)
                tc.append(p)
            else:
                # Para celdas normales, procesar markdown
                _procesar_contenido_celda(val, tc)
            
            tr.append(tc)
        tbl.append(tr)

    add_row(header, bold=True)
    for row in data_rows:
        add_row(row)


# Nuevas funciones auxiliares para identificar tablas
def _is_table_separator_line(line: str) -> bool:
    """
    Detecta línea separadora de tabla markdown (soporta :, ---).
    """
    pattern = r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?\s*$"
    return re.fullmatch(pattern, line) is not None


def _is_table_header_candidate(line: str) -> bool:
    """
    Candidata a header si contiene '|' y no es separador.
    """
    return "|" in line and not _is_table_separator_line(line)


def insertar_md_como_parrafos(para: Paragraph, md_text: str) -> None:
    parent = para._element.getparent()
    idx = parent.index(para._element)
    parent.remove(para._element)

    bloques = md_text.strip().splitlines()
    doc = para._parent
    insert_count = 0
    i = 0

    while i < len(bloques):
        ln = bloques[i]

        # Detección mejorada de tabla
        if _is_table_header_candidate(ln):
            # Saltar líneas en blanco para buscar separador
            j = i + 1
            while j < len(bloques) and bloques[j].strip() == "":
                j += 1
            if j < len(bloques) and _is_table_separator_line(bloques[j]):
                # Recolectar líneas de tabla
                tabla_md = [bloques[i], bloques[j]]
                k = j + 1
                while k < len(bloques):
                    line_k = bloques[k]
                    if not line_k.strip():
                        break
                    if "|" not in line_k:
                        break
                    tabla_md.append(line_k)
                    k += 1
                render_tabla_md(tabla_md, idx + insert_count, parent, doc)
                insert_count += 1
                i = k
                continue

        ln_strip = ln.strip()
        if not ln_strip:
            i += 1
            continue

        new_p = OxmlElement("w:p")
        parent.insert(idx + insert_count, new_p)
        new_para = Paragraph(new_p, doc)
        insert_count += 1

        stripped = ln_strip
        estilo = None
        fuente_manual = None

        if stripped.startswith("##### "):
            stripped = stripped[6:].strip()
            estilo = "Heading 5"
            fuente_manual = Pt(10)
        elif stripped.startswith("#### "):
            stripped = stripped[5:].strip()
            fuente_manual = Pt(11)
        elif stripped.startswith("### "):
            stripped = stripped[4:].strip()
            estilo = "Heading 3"
        elif stripped.startswith("## "):
            stripped = stripped[3:].strip()
            estilo = "Heading 2"
            fuente_manual = Pt(14)
        elif stripped.startswith("- "):
            stripped = stripped[2:].strip()
            try:
                new_para.style = "List Bullet"
            except KeyError:
                # Si no existe el estilo, usar formato manual
                from docx.shared import Inches
                new_para.paragraph_format.left_indent = Inches(0.25)
                new_para.paragraph_format.first_line_indent = Inches(-0.25)
                stripped = "• " + stripped

        if estilo:
            new_para.style = estilo

        parts = re.split(r"(\*\*[^*]+\*\*)", stripped)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = new_para.add_run(part[2:-2])
                run.bold = True
            else:
                run = new_para.add_run(part)
            if fuente_manual:
                run.font.size = fuente_manual

        i += 1


def reemplazo_inline(p: Paragraph, mapa: Dict[str, str]) -> None:
    patron = re.compile(r"\{\{\s*([^}]+?)\s*\}\}")

    original = p.text or ""
    partes = patron.split(original)

    if len(partes) == 1:
        return

    nuevo = ""
    for i, parte in enumerate(partes):
        if i % 2 == 0:
            nuevo += parte
        else:
            clave = normalizer(parte)
            nuevo += mapa.get(clave, f"{{{{{parte}}}}}")

    estilo = p.style
    p.text = nuevo
    p.style = estilo


def rellenar_y_devolver_bytes(plantilla: bytes, secciones_dict: Dict) -> bytes:
    print("Usando plantilla")
    plantilla_buf = BytesIO(plantilla)
    doc = Document(plantilla_buf)

    print("Construyendo mapa normalizado")
    norm_map = {
        normalizer(s["nombre"]): eliminar_linea(s["contenido"])
        for s in secciones_dict.get("secciones", [])
    }

    titulo_doc = secciones_dict.get("titulo_doc", "")
    norm_map[normalizer("titulo")] = titulo_doc
    doc.core_properties.title = titulo_doc

    print("Rellenando plantilla")
    for p in doc.paragraphs:
        m = re.fullmatch(r"\s*\{\{\s*([^\}]+?)\s*\}\}\s*", p.text or "")
        if m:
            key = normalizer(m.group(1).strip())
            if key in norm_map:
                insertar_md_como_parrafos(p, norm_map[key])
        else:
            reemplazo_inline(p, norm_map)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()