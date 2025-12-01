from sqlalchemy.ext.asyncio import AsyncSession
from .repository import ExecutionRepo
from src.modules.section_execution.service import SectionExecutionService
from src.modules.docx_template.service import DocxTemplateService
from src.modules.section.service import SectionService
from src.modules.job.service import JobService
from src.modules.job.models import Job
from src.modules.llm.service import LLMService
from src.modules.search.service import ChunkService
from .models import Execution, Status
from src.modules.section_execution.models import SectionExecution
from src.modules.section.models import Section
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import json
from .utils import insertar_md_como_parrafos, rellenar_y_devolver_bytes


class ExecutionService:
    def __init__(self, session: AsyncSession):
        self.session = session
        self.execution_repo = ExecutionRepo(session)
        self.job_service = JobService(session)
        self.section_service = SectionService(session)
        self.llm_service = LLMService(session)
        self.section_exec_service = SectionExecutionService(session)
        
    
    @staticmethod
    def _generate_docx_no_template(data: dict) -> bytes:
        """
        Generate a Word document without a template.
        """
        doc = Document()

        title = data.get("titulo_doc", "No Title").strip()
        p_title = doc.add_paragraph(title)
        p_title.style = "Title"
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.core_properties.title = title
        doc.add_page_break()

        secciones = data.get("secciones", [])
        for num, sec in enumerate(secciones):
            contenido = (sec.get("contenido") or "").strip()
            if not contenido:
                continue
            placeholder = doc.add_paragraph("{{_md_block_}}")
            insertar_md_como_parrafos(placeholder, contenido)
            if num < len(secciones) - 1:
                doc.add_page_break()

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
    
    
    @staticmethod
    def _generate_docx_with_template(template: bytes, data: dict) -> bytes:
        """
        Generate a Word document using a template.
        """
        try:
            docx_bytes = rellenar_y_devolver_bytes(template, data)
            return docx_bytes
        except Exception as e:
            raise ValueError(f"Error generating document from template: {str(e)}")
    
    
    async def get_execution(self, execution_id: str):
        """
        Retrieve an execution by its ID.
        """
        execution = await self.execution_repo.get_by_id(execution_id)
        if not execution:
            raise ValueError(f"Execution with ID {execution_id} not found.")
        return execution
    
    async def check_execution_exists(self, execution_id: str) -> bool:
        """
        Check if an execution exists by its ID.
        """
        execution = await self.execution_repo.get_by_id(execution_id)
        return execution is not None
    
    async def get_execution_object(self, execution_id: str, with_model: bool = False):
        """
        Retrieve an execution object by its ID.
        """
        execution = await self.execution_repo.get_execution(execution_id, with_model=with_model)
        if not execution:
            raise ValueError(f"Execution with ID {execution_id} not found.")
        return execution
    
    async def get_number_of_executions(self, document_id: str) -> int:
        """
        Get the number of executions for a specific document.
        """
        count = await self.execution_repo.count_executions_by_document_id(document_id)
        return count
        
    
    async def get_sections_by_execution_id(self, execution_id: str) -> list[SectionExecution]:
        """
        Retrieve all sections associated with a specific execution.
        """
        execution = await self.execution_repo.get_execution_sections(execution_id)
        sections = execution.sections_executions
        if not sections:
            return []
        return sections

    async def get_executions_by_doc_id(self, document_id: str) -> list:
        """
        Retrieve all executions.
        """
        executions = await self.execution_repo.get_executions_by_doc_id(document_id)
        if not executions:
            return []
        return executions
    
    async def create_section_execution_bulk(self, sections: list[Section], execution_id: str) -> list[SectionExecution]:
        """
        Create section executions in bulk for the given sections and execution ID.
        """
        print("Creating section executions in bulk...")
        section_executions = []
        for section in sections:
            section_execution = SectionExecution(
                name=section.name,
                section_id=section.id,
                execution_id=execution_id,
                prompt=section.prompt,
                order=section.order
            )
            section_executions.append(section_execution)
        
        created_section_executions = await self.section_exec_service.add_section_executions_bulk(section_executions)
        print("Created", len(created_section_executions), "section executions.")
        return created_section_executions
    
    async def create_execution(self, document_id: str, llm_id: str) -> Execution:
        """
        Create a new execution for a document.
        """
        llm = await self.llm_service.check_llm_exists(llm_id)
        if not llm:
            raise ValueError(f"LLM with ID {llm_id} not found.")
        
        number_of_executions = await self.get_number_of_executions(document_id)
        new_execution = Execution(
            name=f"Version {number_of_executions + 1}",
            document_id=document_id,
            status=Status.PENDING,
            model_id=llm_id
            )
        execution = await self.execution_repo.add(new_execution)
        if not execution:
            raise ValueError(f"Failed to create execution for document ID {document_id}.")
        
        _ = await self.create_section_execution_bulk(
            await self.section_service.get_document_sections_graph(document_id),
            execution.id)
        
        return execution
    
    async def get_execution_status(self, execution_id: str) -> str:
        """
        Get the status of a specific execution.
        """
        execution = await self.execution_repo.get_execution(execution_id)
        if not execution:
            raise ValueError(f"Execution with ID {execution_id} not found.")
        return execution.status
    
    async def delete_execution(self, execution_id: str):
        """
        Delete an execution by its ID.
        """
        execution = await self.execution_repo.get_execution(execution_id)
        if not execution:
            raise ValueError(f"Execution with ID {execution_id} not found.")
        
        await self.execution_repo.delete(execution)
        return {"message": f"Execution with ID {execution_id} deleted successfully."}
    
    async def clone_execution(self, execution_id: str) -> Execution:
        """
        Clone a completed or approved execution along with its section executions.
        """
        execution = await self.execution_repo.get_execution_sections(execution_id)
        if execution.status not in (Status.COMPLETED, Status.APPROVED):
            raise ValueError("Only completed or approved executions can be cloned.")
        
        new_execution = Execution(
            name=f"{execution.name} copy",
            user_instruction=execution.user_instruction,
            status=Status.COMPLETED,
            status_message=execution.status_message,
            document_id=execution.document_id,
            model_id=execution.model_id,
        )
        new_execution = await self.execution_repo.add(new_execution)
        
        section_execution_service = SectionExecutionService(self.session)
        for section_exec in sorted(execution.sections_executions, key=lambda se: se.order):
            cloned_section_exec = SectionExecution(
                name=section_exec.name,
                user_instruction=section_exec.user_instruction,
                prompt=section_exec.prompt,
                order=section_exec.order,
                output=section_exec.output,
                custom_output=section_exec.custom_output,
                is_locked=section_exec.is_locked,
                section_id=section_exec.section_id,
                execution_id=new_execution.id,
            )
            await section_execution_service.add_section_execution(cloned_section_exec)
        
        return new_execution
    
    async def update_llm(self, execution_id: str, llm_id: str):
        """
        Update the LLM used for an execution.
        """
        execution = await self.execution_repo.get_execution(execution_id)
        if not execution:
            raise ValueError(f"Execution with ID {execution_id} not found.")
        
        llm = await LLMService(self.session).check_llm_exists(llm_id)
        if not llm:
            raise ValueError(f"LLM with ID {llm_id} not found.")
        
        execution.model_id = llm.id
        updated_execution = await self.execution_repo.update(execution)
        return updated_execution
    
    async def modify_section_exec_content(self, section_execution_id: str, new_content: str):
        """
        Modify the content of a section execution.
        """
        section_exec_service = SectionExecutionService(self.session)
        updated_section_exec = await section_exec_service.update_section_execution_content(
            section_execution_id, new_content
        )
        return updated_section_exec
    
    # async def export_execution_markdown(self, execution_id: str) -> list:
    #     """
    #     Export the section executions for a specific execution.
    #     """
    #     section_execs = await self.section_exec_repo.get_sections_by_execution_id(execution_id)
    #     if not section_execs:
    #         raise ValueError(f"No section executions found for execution ID {execution_id}.")
    #     print("results", section_execs)
    #     sorted_execs = sorted(
    #         section_execs,
    #         key=lambda x: (x.section.order)
    #     )
    #     export_data = "\n\n-------\n\n".join([i.custom_output if i.custom_output else i.output for i in sorted_execs])
    #     return export_data
    
    async def export_execution_markdown(self, execution_id: str) -> str:
        """
        Export the section executions for a specific execution.
        """
        execution = await self.execution_repo.get_execution_sections(execution_id)
        section_execs = execution.sections_executions
        if not section_execs:
            raise ValueError(f"No section executions found for execution ID {execution_id}.")
        
        sorted_execs = sorted(
            section_execs,
            key=lambda x: (x.order)
        )
        export_data = "\n\n-------\n\n".join([
            i.custom_output if i.custom_output else i.output 
            for i in sorted_execs
        ])
        return export_data
    
    
    
    # async def export_execution_word(self, execution_id: str) -> bytes:
    #     """
    #     Export the results of a specific execution as a downloadable Word file.
    #     """
    #     section_execs = await self.section_exec_repo.get_sections_by_execution_id(execution_id)
    #     if not section_execs:
    #         raise ValueError(f"No section executions found for execution ID {execution_id}.")
        
    #     # Obtener el nombre del documento desde la primera section execution
    #     document_name = section_execs[0].execution.document.name if section_execs[0].execution and section_execs[0].execution.document else "No Title"
        
    #     sorted_execs = sorted(
    #         section_execs,
    #         key=lambda x: (x.section.order)
    #     )
        
    #     data = {
    #         "titulo_doc": document_name,
    #         "secciones": [
    #             {
    #                 "nombre": sec_exec.section.name if sec_exec.section else "No Name",
    #                 "contenido": sec_exec.custom_output if sec_exec.custom_output else sec_exec.output
    #             }
    #             for sec_exec in sorted_execs
    #         ]
    #     }
        
    #     docx_bytes = self._generate_docx_no_template(data)
    #     return docx_bytes
    
    
    async def get_content_to_word(self, execution_id: str) -> dict:
        """
        Get the content data for generating a Word document.
        """
        execution = await self.execution_repo.get_execution_sections(execution_id)
        
        section_execs = execution.sections_executions
        if not section_execs:
            raise ValueError(f"No section executions found for execution ID {execution_id}.")
        
        # Obtener el nombre del documento desde la ejecución
        document_name = execution.document.name if execution and execution.document else "No Title"
        
        sorted_execs = sorted(
            section_execs,
            key=lambda x: (x.order)
        )
        data = {
            "titulo_doc": document_name,
            "secciones": [
                {
                    "nombre": sec_exec.name if sec_exec.name else "No Name",
                    "contenido": sec_exec.custom_output if sec_exec.custom_output else sec_exec.output
                }
                for sec_exec in sorted_execs
            ]
        }
        return data
    
    async def export_execution_word(self, execution_id: str) -> bytes:
        """
        Export the results of a specific execution as a downloadable Word file.
        """
        data = await self.get_content_to_word(execution_id)
        docx_bytes = self._generate_docx_no_template(data)
        return docx_bytes
    
    # async def export_custom_word(self, execution_id: str) -> bytes:
    #     """
    #     Export a custom Word document based on provided data using the document's template.
    #     """
    #     section_execs = await self.section_exec_repo.get_sections_by_execution_id(execution_id)
    #     if not section_execs:
    #         raise ValueError(f"No section executions found for execution ID {execution_id}.")
        
    #     document_id = section_execs[0].execution.document_id
        
    #     docx_template = await self.docx_template_repo.get_by_document_id(document_id)
    #     if not docx_template:
    #         raise ValueError(f"No DOCX template found for document ID {document_id}.")
        
    #     document_name = section_execs[0].execution.document.name if section_execs[0].execution and section_execs[0].execution.document else "No Title"
        
    #     sorted_execs = sorted(
    #         section_execs,
    #         key=lambda x: (x.section.order)
    #     )
        
    #     data = {
    #         "titulo_doc": document_name,
    #         "secciones": [
    #             {
    #                 "nombre": sec_exec.section.name if sec_exec.section else "No Name",
    #                 "contenido": sec_exec.custom_output if sec_exec.custom_output else sec_exec.output
    #             }
    #             for sec_exec in sorted_execs
    #         ]
    #     }
        
    #     docx_bytes = self._generate_docx_with_template(docx_template.file_data, data)
    #     return docx_bytes
    
    async def export_custom_word(self, execution_id: str) -> bytes:
        """
        Export a custom Word document based on provided data using the document's template.
        """
        data = await self.get_content_to_word(execution_id)
        execution = await self.execution_repo.get_execution_sections(execution_id)
        document_id = execution.document_id
        
        docx_template = await DocxTemplateService(self.session).get_by_document_id(document_id)
        
        docx_bytes = self._generate_docx_with_template(docx_template.file_data, data)
        return docx_bytes
        
        
    
    async def approve_execution(self, execution_id: str):
        """
        Check if other executions of the same document are approved,
        if so replace the status to completed.
        Change the status of the execution to APPROVED.
        """
        execution = await self.execution_repo.get_execution(execution_id)
        if not execution:
            raise ValueError(f"Execution with ID {execution_id} not found.")
        
        chunk_service = ChunkService(self.session)
        try:
            await chunk_service.generate_chunks(execution_id)
        except Exception as e:
            raise ValueError(f"Failed to generate chunks for execution ID {execution_id}: {str(e)}")
        
        # Check if there are other approved executions for the same document
        past_approved_execution = await self.execution_repo.get_approved_execution_by_doc_id(execution.document_id)
        if past_approved_execution:
            past_approved_execution.status = Status.COMPLETED
            await self.execution_repo.update(past_approved_execution)
            await chunk_service.delete_chunks_by_execution(past_approved_execution.id)
            
        execution.status = Status.APPROVED
        updated_execution = await self.execution_repo.update(execution)
        return updated_execution
    
    
    async def disapprove_execution(self, execution_id: str):
        """
        Change the status of the execution to DISAPPROVED.
        """
        execution = await self.execution_repo.get_execution(execution_id)
        if not execution:
            raise ValueError(f"Execution with ID {execution_id} not found.")
        
        chunk_service = ChunkService(self.session)
        try:
            await chunk_service.delete_chunks_by_execution(execution_id)
        except Exception as e:
            raise ValueError(f"Failed to delete chunks for execution ID {execution_id}: {str(e)}")
        
        execution.status = Status.COMPLETED
        updated_execution = await self.execution_repo.update(execution)
        return updated_execution
    
    async def get_execution_for_chunking(self, execution_id: str) -> Execution:
        """
        Retrieve an execution for chunking purposes.
        """
        execution = await self.execution_repo.get_execution_to_chunking(execution_id)
        if not execution:
            raise ValueError(f"Execution with ID {execution_id} not found.")
        return execution
    
    async def update_status(self, execution_id: str, status: Status, status_message: str, user_instructions: str = None) -> Execution:
        """
        Update the status of an execution.
        """
        execution = await self.execution_repo.get_execution(execution_id)
        if not execution:
            raise ValueError(f"Execution with ID {execution_id} not found.")
        
        execution.status = status
        execution.status_message = status_message
        if user_instructions is not None:
            execution.user_instructions = user_instructions
        
        updated_execution = await self.execution_repo.update(execution)
        return updated_execution
    
    async def add_execution_graph_job(self, document_id: str, llm_id: str, execution_id: str = None,
                                      user_instructions: str = None, start_section_id: str = None,
                                      single_section_mode: bool = False) -> Job:
        """
        Enqueue a job to run the generation graph for a document execution.
        """
        
        # Check if document exists
        document = await self.execution_repo.check_if_document_exists(document_id)
        if not document:
            raise ValueError(f"Document with ID {document_id} does not exist.")
        
        # If start_section_id is provided, check if it exists in the document
        if start_section_id:
            section_exists = await self.section_service.check_section_exists(start_section_id, document_id)
            if not section_exists:
                raise ValueError(f"Section with ID {start_section_id} does not exist in document {document_id}.")
            
        # If execution_id is not provided, create a new execution
        if not execution_id:
            execution = await self.create_execution(document_id, llm_id)
            execution_id = str(execution.id)
        else:
            execution = await self.get_execution(execution_id)
            if not execution:
                raise ValueError(f"Execution with ID {execution_id} does not exist.")
            
            
        _ = await self.update_status(execution_id, Status.PENDING, "Pending", user_instructions)
        payload = {
            "document_id": document_id,
            "execution_id": execution_id,
            "llm_id": llm_id,
            "user_instructions": user_instructions,
            "start_section_id": start_section_id,
            "single_section_mode": single_section_mode
        }
        job = await self.job_service.enqueue_job(
            job_type="run_generation_graph",
            payload=json.dumps(payload)
        )
        
        execution_data = await self.get_execution(execution_id)
        data = {
            "execution": execution_data,
            "job": job
        }
        return data
