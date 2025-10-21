from sqlalchemy.ext.asyncio import AsyncSession
from src.database.repositories.execution_repo import ExecutionRepo
from src.database.repositories.sectionexec_repo import SectionExecRepo
from src.services.chunk_service import ChunkService
from src.database.repositories.llm_repo import LLMRepo
from src.database.repositories.docx_template_repo import DocxTemplateRepo
from src.database.models import Execution, Status
from src.config import system_config
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from src.services.docx_utils import insertar_md_como_parrafos, rellenar_y_devolver_bytes


class ExecutionService:
    def __init__(self, session: AsyncSession):
        self.session = session
        self.execution_repo = ExecutionRepo(session)
        self.section_exec_repo = SectionExecRepo(session)
        self.llm_repo = LLMRepo(session)
        self.docx_template_repo = DocxTemplateRepo(session)
        
    
    @staticmethod
    def _generate_docx_no_template(data: dict) -> bytes:
        """
        Generate a Word document without a template.
        """
        doc = Document()

        title = data.get("titulo_doc", "No Title").strip()
        p_title = doc.add_paragraph(title)
        p_title.style = "Title"
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.core_properties.title = title
        doc.add_page_break()

        secciones = data.get("secciones", [])
        for num, sec in enumerate(secciones):
            contenido = (sec.get("contenido") or "").strip()
            if not contenido:
                continue
            placeholder = doc.add_paragraph("{{_md_block_}}")
            insertar_md_como_parrafos(placeholder, contenido)
            if num < len(secciones) - 1:
                doc.add_page_break()

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
    
    
    @staticmethod
    def _generate_docx_with_template(template: bytes, data: dict) -> bytes:
        """
        Generate a Word document using a template.
        """
        try:
            docx_bytes = rellenar_y_devolver_bytes(template, data)
            return docx_bytes
        except Exception as e:
            raise ValueError(f"Error generating document from template: {str(e)}")
    
    
    async def get_execution(self, execution_id: str):
        """
        Retrieve an execution by its ID.
        """
        execution = await self.execution_repo.get_by_id(execution_id)
        if not execution:
            raise ValueError(f"Execution with ID {execution_id} not found.")
        return execution

    async def get_executions_by_doc_id(self, document_id: str) -> list:
        """
        Retrieve all executions.
        """
        executions = await self.execution_repo.get_executions_by_doc_id(document_id)
        if not executions:
            return []
        return executions
    
    async def create_execution(self, document_id: str):
        """
        Create a new execution for a document.
        """
        
        llm_name = system_config.DEFAULT_LLM
        
        llm = await self.llm_repo.get_by_name(llm_name)
        if not llm:
            raise ValueError(f"LLM with name {llm_name} not found.")      
        
        new_execution = Execution(
            document_id=document_id,
            status=Status.PENDING,
            model_id=llm.id
            )
        execution = await self.execution_repo.add(new_execution)
        if not execution:
            raise ValueError(f"Failed to create execution for document ID {document_id}.")
        return execution
    
    async def get_execution_status(self, execution_id: str) -> str:
        """
        Get the status of a specific execution.
        """
        execution = await self.execution_repo.get_by_id(execution_id)
        if not execution:
            raise ValueError(f"Execution with ID {execution_id} not found.")
        return execution.status
    
    async def delete_execution(self, execution_id: str):
        """
        Delete an execution by its ID.
        """
        execution = await self.execution_repo.get_execution(execution_id)
        if not execution:
            raise ValueError(f"Execution with ID {execution_id} not found.")
        
        await self.execution_repo.delete(execution)
        return {"message": f"Execution with ID {execution_id} deleted successfully."}
    
    async def update_llm(self, execution_id: str, llm_id: str):
        """
        Update the LLM used for an execution.
        """
        execution = await self.execution_repo.get_execution(execution_id)
        if not execution:
            raise ValueError(f"Execution with ID {execution_id} not found.")
        
        llm = await self.llm_repo.get_by_id(llm_id)
        if not llm:
            raise ValueError(f"LLM with ID {llm_id} not found.")
        
        execution.model_id = llm.id
        updated_execution = await self.execution_repo.update(execution)
        return updated_execution
    
    async def modify_section_exec_content(self, section_execution_id: str, new_content: str):
        """
        Modify the content of a section execution.
        """
        section_exec = await self.section_exec_repo.get_by_id(section_execution_id)
        if not section_exec:
            raise ValueError(f"Section execution with ID {section_execution_id} not found.")
        
        section_exec.custom_output = new_content
        updated_section_exec = await self.section_exec_repo.update(section_exec)
        return updated_section_exec
    
    async def export_execution_markdown(self, execution_id: str) -> list:
        """
        Export the section executions for a specific execution.
        """
        section_execs = await self.section_exec_repo.get_sections_by_execution_id(execution_id)
        if not section_execs:
            raise ValueError(f"No section executions found for execution ID {execution_id}.")
        print("results", section_execs)
        sorted_execs = sorted(
            section_execs,
            key=lambda x: (x.section.order)
        )
        export_data = "\n\n-------\n\n".join([i.custom_output if i.custom_output else i.output for i in sorted_execs])
        return export_data
    
    async def export_execution_word(self, execution_id: str) -> bytes:
        """
        Export the results of a specific execution as a downloadable Word file.
        """
        section_execs = await self.section_exec_repo.get_sections_by_execution_id(execution_id)
        if not section_execs:
            raise ValueError(f"No section executions found for execution ID {execution_id}.")
        
        # Obtener el nombre del documento desde la primera section execution
        document_name = section_execs[0].execution.document.name if section_execs[0].execution and section_execs[0].execution.document else "No Title"
        
        sorted_execs = sorted(
            section_execs,
            key=lambda x: (x.section.order)
        )
        
        data = {
            "titulo_doc": document_name,
            "secciones": [
                {
                    "nombre": sec_exec.section.name if sec_exec.section else "No Name",
                    "contenido": sec_exec.custom_output if sec_exec.custom_output else sec_exec.output
                }
                for sec_exec in sorted_execs
            ]
        }
        
        docx_bytes = self._generate_docx_no_template(data)
        return docx_bytes
    
    async def export_custom_word(self, execution_id: str) -> bytes:
        """
        Export a custom Word document based on provided data using the document's template.
        """
        section_execs = await self.section_exec_repo.get_sections_by_execution_id(execution_id)
        if not section_execs:
            raise ValueError(f"No section executions found for execution ID {execution_id}.")
        
        document_id = section_execs[0].execution.document_id
        
        docx_template = await self.docx_template_repo.get_by_document_id(document_id)
        if not docx_template:
            raise ValueError(f"No DOCX template found for document ID {document_id}.")
        
        document_name = section_execs[0].execution.document.name if section_execs[0].execution and section_execs[0].execution.document else "No Title"
        
        sorted_execs = sorted(
            section_execs,
            key=lambda x: (x.section.order)
        )
        
        data = {
            "titulo_doc": document_name,
            "secciones": [
                {
                    "nombre": sec_exec.section.name if sec_exec.section else "No Name",
                    "contenido": sec_exec.custom_output if sec_exec.custom_output else sec_exec.output
                }
                for sec_exec in sorted_execs
            ]
        }
        
        docx_bytes = self._generate_docx_with_template(docx_template.file_data, data)
        return docx_bytes
    
    async def approve_execution(self, execution_id: str):
        """
        Check if other executions of the same document are approved,
        if so replace the status to completed.
        Change the status of the execution to APPROVED.
        """
        execution = await self.execution_repo.get_execution(execution_id)
        if not execution:
            raise ValueError(f"Execution with ID {execution_id} not found.")
        
        chunk_service = ChunkService(self.session)
        try:
            await chunk_service.generate_chunks(execution_id)
        except Exception as e:
            raise ValueError(f"Failed to generate chunks for execution ID {execution_id}: {str(e)}")
        
        # Check if there are other approved executions for the same document
        past_approved_execution = await self.execution_repo.get_approved_execution_by_doc_id(execution.document_id)
        if past_approved_execution:
            past_approved_execution.status = Status.COMPLETED
            await self.execution_repo.update(past_approved_execution)
            await chunk_service.delete_chunks_by_execution(past_approved_execution.id)
            
        execution.status = Status.APPROVED
        updated_execution = await self.execution_repo.update(execution)
        return updated_execution
    
    
    async def disapprove_execution(self, execution_id: str):
        """
        Change the status of the execution to DISAPPROVED.
        """
        execution = await self.execution_repo.get_execution(execution_id)
        if not execution:
            raise ValueError(f"Execution with ID {execution_id} not found.")
        
        chunk_service = ChunkService(self.session)
        try:
            await chunk_service.delete_chunks_by_execution(execution_id)
        except Exception as e:
            raise ValueError(f"Failed to delete chunks for execution ID {execution_id}: {str(e)}")
        
        execution.status = Status.COMPLETED
        updated_execution = await self.execution_repo.update(execution)
        return updated_execution
