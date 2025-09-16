import re
import unicodedata
from typing import Dict, List
from io import BytesIO
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Pt


def normalizer(s: str) -> str:
    s = unicodedata.normalize("NFKD", s)
    s = "".join(c for c in s if not unicodedata.combining(c))
    s = s.lower()
    s = re.sub(r"[^a-z0-9]+", " ", s).strip()
    s = re.sub(r"\s+", " ", s)
    return s


def eliminar_linea(md: str) -> str:
    out = []
    for line in md.splitlines():
        if re.fullmatch(r"\s*-{3,}\s*", line):
            out.append("")
        else:
            out.append(line)
    return "\n".join(out)


def extraer_secciones_placeholders(archivo_md: str) -> Dict:
    with open(archivo_md, "r", encoding="utf-8") as f:
        texto = f.read()

    titulo_doc = None
    for ln in texto.splitlines():
        if ln.startswith("# "):
            titulo_doc = ln.replace("#", "").strip()
            break

    token_line = re.compile(r"(?m)^\s*\{\{\s*([^\}]+?)\s*\}\}\s*$")
    matches = list(token_line.finditer(texto))

    secciones: List[Dict] = []
    for i, m in enumerate(matches):
        key = m.group(1).strip()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(texto)
        contenido = texto[start:end].lstrip()
        contenido = eliminar_linea(contenido)
        secciones.append({"titulo": key, "contenido": contenido})

    return {"titulo_doc": titulo_doc, "secciones": secciones}


def render_tabla_md(md_lines: List[str], insert_idx: int, parent_element, doc) -> None:
    if len(md_lines) < 2:
        return

    header = [h.strip() for h in md_lines[0].split("|") if h.strip()]
    separator = [s.strip() for s in md_lines[1].split("|") if s.strip()]

    if len(header) != len(separator):
        return

    tbl = OxmlElement("w:tbl")
    parent_element.insert(insert_idx, tbl)

    def add_row(values):
        tr = OxmlElement("w:tr")
        for val in values:
            tc = OxmlElement("w:tc")
            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.text = val
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    add_row(header)
    for line in md_lines[2:]:
        cells = [c.strip() for c in line.split("|") if c.strip()]
        if len(cells) == len(header):
            add_row(cells)


def insertar_md_como_parrafos(para: Paragraph, md_text: str) -> None:
    parent = para._element.getparent()
    idx = parent.index(para._element)
    parent.remove(para._element)

    bloques = md_text.strip().splitlines()
    doc = para._parent
    insert_count = 0

    i = 0
    while i < len(bloques):
        ln = bloques[i].strip()

        if i + 1 < len(bloques) and re.fullmatch(r"\s*\|?(\s*-+\s*\|)+\s*", bloques[i + 1]):
            j = i
            tabla_md = []
            while j < len(bloques) and "|" in bloques[j]:
                tabla_md.append(bloques[j])
                j += 1
            
            render_tabla_md(tabla_md, idx + insert_count, parent, doc)

            insert_count += 1
            i = j
            continue

        if not ln:
            i += 1
            continue

        new_p = OxmlElement("w:p")
        parent.insert(idx + insert_count, new_p)
        new_para = Paragraph(new_p, doc)
        insert_count += 1

        stripped = ln.strip()
        estilo = None
        fuente_manual = None

        if stripped.startswith("#### "):
            stripped = stripped[5:].strip()
            fuente_manual = Pt(11)

        elif stripped.startswith("### "):
            stripped = stripped[4:].strip()
            estilo = "Heading 3"

        elif stripped.startswith("## "):
            stripped = stripped[3:].strip()
            estilo = "Heading 2"
            fuente_manual = Pt(14)

        elif stripped.startswith("- "):
            stripped = stripped[2:].strip()
            estilo = "List Bullet"

        if estilo:
            new_para.style = estilo

        parts = re.split(r"(\*\*[^*]+\*\*)", stripped)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = new_para.add_run(part[2:-2])
                run.bold = True
            else:
                run = new_para.add_run(part)

            if fuente_manual:
                run.font.size = fuente_manual

        i += 1


def reemplazo_inline(p: Paragraph, mapa: Dict[str, str]) -> None:
    patron = re.compile(r"\{\{\s*([^}]+?)\s*\}\}")

    original = p.text or ""
    partes = patron.split(original)

    if len(partes) == 1:
        return

    nuevo = ""
    for i, parte in enumerate(partes):
        if i % 2 == 0:
            nuevo += parte
        else:
            clave = normalizer(parte)
            nuevo += mapa.get(clave, f"{{{{{parte}}}}}")

    estilo = p.style
    p.text = nuevo
    p.style = estilo


def rellenar_y_devolver_bytes(plantilla_path: str, secciones_dict: Dict) -> bytes:
    doc = Document(plantilla_path)

    norm_map = {
        normalizer(s["titulo"]): eliminar_linea(s["contenido"])
        for s in secciones_dict.get("secciones", [])
    }

    titulo_doc = secciones_dict.get("titulo_doc", "")
    norm_map[normalizer("titulo")] = titulo_doc
    doc.core_properties.title = titulo_doc

    for p in doc.paragraphs:
        m = re.fullmatch(r"\s*\{\{\s*([^\}]+?)\s*\}\}\s*", p.text or "")
        if m:
            key = normalizer(m.group(1).strip())
            if key in norm_map:
                insertar_md_como_parrafos(p, norm_map[key])
        else:
            reemplazo_inline(p, norm_map)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()